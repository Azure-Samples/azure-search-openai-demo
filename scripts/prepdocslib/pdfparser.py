import docx
import html
import os
from abc import ABC
from typing import IO, AsyncGenerator, Union

from azure.ai.formrecognizer import DocumentTable
from azure.ai.formrecognizer.aio import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
from azure.core.credentials_async import AsyncTokenCredential
from docx import Document
from pypdf import PdfReader

from .strategy import USER_AGENT


class Page:
    """
    A single page from a pdf

    Attributes:
        page_num (int): Page number
        offset (int): If the text of the entire PDF was concatenated into a single string, the index of the first character on the page. For example, if page 1 had the text "hello" and page 2 had the text "world", the offset of page 2 is 5 ("hellow")
        text (str): The text of the page
    """

    def __init__(self, page_num: int, offset: int, text: str):
        self.page_num = page_num
        self.offset = offset
        self.text = text


class PdfParser(ABC):
    """
    Abstract parser that parses PDFs into pages
    """

    async def parse(self, content: IO) -> AsyncGenerator[Page, None]:
        if False:
            yield


class LocalPdfParser(PdfParser):
    """
    Concrete parser backed by PyPDF that can parse PDFs into pages
    To learn more, please visit https://pypi.org/project/pypdf/
    """

    async def parse(self, content: IO) -> AsyncGenerator[Page, None]:
        reader = PdfReader(content)
        pages = reader.pages
        offset = 0
        for page_num, p in enumerate(pages):
            page_text = p.extract_text()
            yield Page(page_num=page_num, offset=offset, text=page_text)
            offset += len(page_text)


class DocumentAnalysisPdfParser(PdfParser):
    """
    Concrete parser backed by Azure AI Document Intelligence that can parse PDFS into pages
    To learn more, please visit https://learn.microsoft.com/azure/ai-services/document-intelligence/overview
    """

    def __init__(
        self,
        endpoint: str,
        credential: Union[AsyncTokenCredential, AzureKeyCredential],
        model_id="prebuilt-layout",
        verbose: bool = False,
    ):
        self.model_id = model_id
        self.endpoint = endpoint
        self.credential = credential
        self.verbose = verbose

    async def parse(self, content: IO) -> AsyncGenerator[Page, None]:
        if self.verbose:
            print(f"Extracting text from '{content.name}' using Azure Document Intelligence")

        async with DocumentAnalysisClient(
            endpoint=self.endpoint, credential=self.credential, headers={"x-ms-useragent": USER_AGENT}
        ) as form_recognizer_client:
            
            # Parse PDF
            if os.path.splitext(content.name)[1].lower() == ".pdf":
                poller = await form_recognizer_client.begin_analyze_document(model_id=self.model_id, document=content)
                form_recognizer_results = await poller.result()

                #JHOLLASDELETE
                with open('./out3.txt', 'a') as file:
                        file.write(''.join(str(table) for table in form_recognizer_results.tables))
                        file.write(''.join(str(content) for content in form_recognizer_results.content))

                offset = 0
                for page_num, page in enumerate(form_recognizer_results.pages):
                    tables_on_page = [
                        table
                        for table in (form_recognizer_results.tables or [])
                        if table.bounding_regions and table.bounding_regions[0].page_number == page_num + 1
                    ]

                    # mark all positions of the table spans in the page
                    page_offset = page.spans[0].offset
                    page_length = page.spans[0].length
                    table_chars = [-1] * page_length
                    for table_id, table in enumerate(tables_on_page):
                        for span in table.spans:
                            # replace all table spans with "table_id" in table_chars array
                            for i in range(span.length):
                                idx = span.offset - page_offset + i
                                if idx >= 0 and idx < page_length:
                                    table_chars[idx] = table_id

                    # build page text by replacing characters in table spans with table html
                    page_text = ""
                    added_tables = set()
                    for idx, table_id in enumerate(table_chars):
                        if table_id == -1:
                            page_text += form_recognizer_results.content[page_offset + idx]
                        elif table_id not in added_tables:
                            page_text += DocumentAnalysisPdfParser.pdf_table_to_html(tables_on_page[table_id])
                            added_tables.add(table_id)

                    #JHOLLASDELETE
                    with open('./out1.txt', 'a') as file:
                        file.write(page_text)

                    yield Page(page_num=page_num, offset=offset, text=page_text)
                    offset += len(page_text)

            # Parse docx
            elif os.path.splitext(content.name)[1].lower() == ".docx":
                doc = docx.Document(content.name)
                page_text = ""
                i = 0

                for element in doc.element.body:
                    if element.tag.endswith('tbl'):
                        # Extract table data
                        table_data = []
                        cell_ids = []
                        for row in doc.tables[i].rows:
                            row_data = []
                            for cell in row.cells:
                                cell_id = str(cell._tc.left)+str(cell._tc.right)+str(cell._tc.top)+str(cell._tc.bottom)
                                if(cell_id not in cell_ids):
                                    cell_ids.append(cell_id)
                                    cell_data = {"row_span":cell._tc.bottom-cell._tc.top,"column_span":cell._tc.grid_span,"content":cell.text.replace("\n", "").replace("\r", "")}
                                    row_data.append(cell_data)
                            table_data.append(row_data)
                        page_text += DocumentAnalysisPdfParser.docx_table_to_html(table_data)
                        i += 1
                    elif element.tag.endswith('p'):
                        # Extract paragraph text
                        paragraph_text = element.text
                        page_text += paragraph_text + '\n'

                #JHOLLASDELETE
                with open('./out2.txt', 'a') as file:
                        file.write(page_text)
                yield Page(page_num=1, offset=0, text=page_text)

    @classmethod
    def pdf_table_to_html(cls, table: DocumentTable):
        table_html = "<table>"
        rows = [
            sorted([cell for cell in table.cells if cell.row_index == i], key=lambda cell: cell.column_index)
            for i in range(table.row_count)
        ]
        for row_cells in rows:
            table_html += "<tr>"
            for cell in row_cells:
                tag = "th" if (cell.kind == "columnHeader" or cell.kind == "rowHeader") else "td"
                cell_spans = ""
                if cell.column_span is not None and cell.column_span > 1:
                    cell_spans += f" colSpan={cell.column_span}"
                if cell.row_span is not None and cell.row_span > 1:
                    cell_spans += f" rowSpan={cell.row_span}"
                table_html += f"<{tag}{cell_spans}>{html.escape(cell.content)}</{tag}>"
            table_html += "</tr>"
        table_html += "</table>"
        return table_html

    @staticmethod
    def docx_table_to_html(table):
        table_html = "<table>"

        for rows in table:
            table_html += "<tr>"
            for cell in rows:
                span = ""
                test = cell['row_span']
                if cell['row_span'] > 1:
                    span = f" rowspan=\"{cell['row_span']}\""
                if cell['column_span']>1:
                    span += f" colspan=\"{cell['column_span']}\""
                table_html += f"<td{span}>{cell['content']}</td>"
            table_html += "</tr>"
        table_html += "</table>\n"
        return table_html